#!/usr/bin/env python3
"""
build-docs.py — regenerate hardgate's multi-format docs from markdown.

Produces:
    README.txt, README.docx, README.pdf        (from README.md)
    USER-MANUAL.docx, USER-MANUAL.pdf          (from USER-MANUAL.md)

Both .docx and .pdf outputs embed the UML architecture diagram from
docs/architecture.mmd as a source-code appendix.

Requires: python-docx, reportlab
    python3 -m pip install --user python-docx reportlab
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Preformatted,
        PageBreak,
    )
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib import colors
except ImportError as e:
    sys.stderr.write(
        f"ERROR: missing dependency ({e.name}). Install with:\n"
        "    python3 -m pip install --user python-docx reportlab\n"
    )
    sys.exit(1)

REPO_ROOT = Path(__file__).resolve().parent.parent
COPYRIGHT = "Copyright (c) 2026 Scott Converse - MIT License"


def md_to_txt(md: str) -> str:
    """Convert markdown to plain text (strip fences, keep code indented)."""
    out: list[str] = []
    in_code = False
    for line in md.splitlines():
        if line.startswith("```"):
            in_code = not in_code
            out.append("")
            continue
        if in_code:
            out.append("    " + line)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            out.append("")
            out.append(text)
            if level == 1:
                out.append("=" * min(len(text), 72))
            elif level == 2:
                out.append("-" * min(len(text), 72))
            continue
        t = re.sub(r"`([^`]+)`", r"\1", line)
        t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
        t = re.sub(r"\*([^*]+)\*", r"\1", t)
        t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", t)
        out.append(t)
    return "\n".join(out)


def md_to_docx(
    md: str, out_path: Path, title: str, include_uml: bool = False
) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(title, level=0)
    p = doc.add_paragraph(COPYRIGHT)
    for r in p.runs:
        r.italic = True

    in_code = False
    code_buf: list[str] = []
    for line in md.splitlines():
        if line.startswith("```"):
            if in_code:
                cp = doc.add_paragraph()
                run = cp.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            continue
        stripped = line.strip()
        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue
        if stripped:
            t = re.sub(r"`([^`]+)`", r"\1", line)
            t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
            t = re.sub(r"\*([^*]+)\*", r"\1", t)
            t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", t)
            doc.add_paragraph(t)
        else:
            doc.add_paragraph("")

    if include_uml:
        doc.add_page_break()
        doc.add_heading(
            "Appendix: Architecture Diagram (UML - Mermaid source)", level=1
        )
        doc.add_paragraph(
            "The following Mermaid diagram describes the 7-location "
            "enforcement model. Render it at https://mermaid.live or with "
            "any Mermaid-capable tool."
        )
        uml = (REPO_ROOT / "docs" / "architecture.mmd").read_text(encoding="utf-8")
        p = doc.add_paragraph()
        run = p.add_run(uml)
        run.font.name = "Consolas"
        run.font.size = Pt(8)

    doc.save(str(out_path))


def md_to_pdf(
    md: str, out_path: Path, title: str, include_uml: bool = False
) -> None:
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title=title,
        author="Scott Converse",
    )
    ss = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "h1", parent=ss["Heading1"], fontSize=20,
        textColor=colors.HexColor("#b02020"), spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "h2", parent=ss["Heading2"], fontSize=14,
        textColor=colors.HexColor("#333"), spaceBefore=12, spaceAfter=6,
    )
    h3 = ParagraphStyle(
        "h3", parent=ss["Heading3"], fontSize=11,
        textColor=colors.HexColor("#555"), spaceBefore=8, spaceAfter=4,
    )
    body = ParagraphStyle(
        "body", parent=ss["BodyText"], fontSize=10,
        leading=14, spaceAfter=6, alignment=TA_LEFT,
    )
    code = ParagraphStyle(
        "code", parent=ss["Code"], fontSize=8, leading=10,
        backColor=colors.HexColor("#f3f3f3"), leftIndent=8, rightIndent=8,
    )
    small = ParagraphStyle(
        "small", parent=body, fontSize=9, textColor=colors.grey
    )

    story = [
        Paragraph(title, h1),
        Paragraph(COPYRIGHT, small),
        Spacer(1, 12),
    ]

    in_code = False
    code_buf: list[str] = []
    for line in md.splitlines():
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_buf), code))
                story.append(Spacer(1, 6))
                code_buf = []
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            if level == 1:
                story.append(Paragraph(text, h1))
            elif level == 2:
                story.append(Paragraph(text, h2))
            else:
                story.append(Paragraph(text, h3))
            continue
        if line.strip():
            t = re.sub(r"`([^`]+)`", r'<font face="Courier">\1</font>', line)
            t = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", t)
            t = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", t)
            t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<link href="\2">\1</link>', t)
            try:
                story.append(Paragraph(t, body))
            except Exception:
                story.append(Paragraph(line, body))
        else:
            story.append(Spacer(1, 4))

    if include_uml:
        story.append(PageBreak())
        story.append(
            Paragraph("Appendix: Architecture Diagram (UML - Mermaid source)", h2)
        )
        story.append(Paragraph("Render at https://mermaid.live", small))
        story.append(Spacer(1, 8))
        uml = (REPO_ROOT / "docs" / "architecture.mmd").read_text(encoding="utf-8")
        story.append(Preformatted(uml, code))

    doc.build(story)


def main() -> int:
    readme_md = (REPO_ROOT / "README.md").read_text(encoding="utf-8")
    manual_md = (REPO_ROOT / "USER-MANUAL.md").read_text(encoding="utf-8")

    print("Regenerating multi-format docs from markdown")
    print("=" * 50)

    outputs = [
        ("README.txt",       lambda: (REPO_ROOT / "README.txt").write_text(md_to_txt(readme_md), encoding="utf-8")),
        ("README.docx",      lambda: md_to_docx(readme_md, REPO_ROOT / "README.docx", "hardgate - README", include_uml=True)),
        ("README.pdf",       lambda: md_to_pdf(readme_md, REPO_ROOT / "README.pdf", "hardgate - README", include_uml=True)),
        ("USER-MANUAL.docx", lambda: md_to_docx(manual_md, REPO_ROOT / "USER-MANUAL.docx", "hardgate - User Manual", include_uml=True)),
        ("USER-MANUAL.pdf",  lambda: md_to_pdf(manual_md, REPO_ROOT / "USER-MANUAL.pdf", "hardgate - User Manual", include_uml=True)),
    ]

    for name, fn in outputs:
        fn()
        size = (REPO_ROOT / name).stat().st_size
        print(f"  wrote {name:20} {size:>8} bytes")

    print("\nDone.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
